from rword.core.docx_io import load_docx, save_docx


def test_save_creates_docx(editor, tmp_path):
    editor.setHtml("<p>Contenido <b>negrita</b></p>")
    path = tmp_path / "doc.docx"
    save_docx(editor, path)
    assert path.exists()
    assert path.read_bytes().startswith(b"PK")


def test_load_docx_text(editor, tmp_path):
    editor.setHtml("<p>Hola <b>negrita</b> y <i>cursiva</i></p><p>Segundo párrafo</p>")
    path = tmp_path / "doc.docx"
    save_docx(editor, path)

    editor.clear()
    load_docx(editor, path)
    text = editor.toPlainText()
    assert "Hola" in text
    assert "negrita" in text
    assert "cursiva" in text
    assert "Segundo párrafo" in text


def test_load_docx_preserves_bold(editor, tmp_path):
    editor.setHtml("<p>Esto es <b>muy importante</b>.</p>")
    path = tmp_path / "doc.docx"
    save_docx(editor, path)

    editor.clear()
    load_docx(editor, path)
    html = editor.toHtml()
    assert "<b" in html or "font-weight" in html


def test_editor_load_save_docx(editor, tmp_path):
    editor.setHtml("<p>Guardar y abrir con el editor.</p>")
    path = tmp_path / "doc.docx"
    editor.save_file(path)
    assert path.exists()

    editor.clear()
    editor.set_file_path(None)
    editor.load_file(path)
    assert "Guardar y abrir" in editor.toPlainText()
    assert editor.file_path == path
    assert not editor.document().isModified()


def test_docx_list_roundtrip(editor, tmp_path):
    editor.setHtml("<ul><li>Uno</li><li>Dos</li></ul>")
    path = tmp_path / "lista.docx"
    save_docx(editor, path)

    editor.clear()
    load_docx(editor, path)
    assert "Uno" in editor.toPlainText()
    assert "Dos" in editor.toPlainText()


def test_docx_table_roundtrip(editor, tmp_path):
    editor.setHtml(
        "<table border='1'><tr><td>A1</td><td>B1</td></tr>"
        "<tr><td>A2</td><td>B2</td></tr></table>"
    )
    path = tmp_path / "tabla.docx"
    save_docx(editor, path)

    editor.clear()
    load_docx(editor, path)
    assert "A1" in editor.toPlainText()
    assert "B2" in editor.toPlainText()


def test_docx_default_format(main_window):
    assert main_window._suggested_name().endswith(".docx")


def test_docx_first_filter(main_window):
    from rword.ui.main_window import FILE_DIALOG_FILTER

    assert FILE_DIALOG_FILTER.startswith("Documento de Word (*.docx)")


def test_save_as_adds_docx_extension(main_window, tmp_path, monkeypatch):
    main_window._editor.setHtml("<p>texto</p>")
    target = tmp_path / "salida"

    def fake_dialog(parent, title, default, file_filter):
        return str(target), "Documento de Word (*.docx)"

    monkeypatch.setattr(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName", fake_dialog
    )
    main_window._save_document_as()
    assert main_window._editor.file_path == tmp_path / "salida.docx"


def test_docx_roundtrip_preserves_full_formatting(editor, tmp_path):
    editor.setHtml(
        "<p style='text-align:center'><b>Centrado y negrita</b></p>"
        "<p><span style='color:#ff0000'>Texto en rojo</span></p>"
    )
    path = tmp_path / "full.docx"
    save_docx(editor, path)

    editor.clear()
    load_docx(editor, path)
    text = editor.toPlainText()
    assert "Centrado y negrita" in text
    assert "Texto en rojo" in text
    html = editor.toHtml()
    assert "ff0000" in html or "#ff0000" in html
    assert "<b" in html or "font-weight" in html


def test_external_docx_preserves_formatting(editor, tmp_path):
    """Un .docx creado por otro editor conserva el formato al abrirlo."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    path = tmp_path / "externo.docx"
    doc = Document()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run("Negrita ")
    run1.bold = True
    run2 = para.add_run("Rojo ")
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run3 = para.add_run("Grande")
    run3.font.size = Pt(20)
    doc.save(str(path))

    from rword.core.docx_io import _read_embedded_html, load_docx

    assert _read_embedded_html(path) is None
    load_docx(editor, path)
    html = editor.toHtml()
    assert "<b" in html or "font-weight" in html
    assert "#ff0000" in html
    assert "center" in html
    assert "20pt" in html


def test_external_docx_image(editor, tmp_path):
    from docx import Document
    from PySide6.QtGui import QColor, QImage

    img_path = tmp_path / "img.png"
    img = QImage(30, 20, QImage.Format.Format_RGB32)
    img.fill(QColor("blue"))
    img.save(str(img_path))

    path = tmp_path / "con_imagen.docx"
    doc = Document()
    doc.add_paragraph("texto")
    doc.add_paragraph().add_run().add_picture(str(img_path))
    doc.save(str(path))

    from rword.core.docx_io import _read_embedded_html, load_docx

    assert _read_embedded_html(path) is None
    editor.clear()
    load_docx(editor, path)
    assert "<img" in editor.toHtml() or "src=" in editor.toHtml()


def test_external_docx_default_font_size(editor, tmp_path):
    from docx import Document

    path = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(str(path))

    from rword.core.docx_io import load_docx

    editor.clear()
    load_docx(editor, path)
    assert editor.document().defaultFont().pointSizeF() == 11.0
