"""Lectura y escritura de documentos en formato .docx (Word)."""

from __future__ import annotations

import html as html_module
from pathlib import Path

from PySide6.QtCore import QUrl
from PySide6.QtGui import QTextDocument, QTextListFormat, QTextTable
from PySide6.QtWidgets import QTextEdit


def _image_html(document, relationship_id: str) -> str:
    """Devuelve <img src='file:///...'> para una imagen incrustada."""
    import uuid

    try:
        image = document.part.related_parts[relationship_id]
    except KeyError:
        return ""
    temp_dir = Path.home() / ".cache" / "rword"
    temp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = temp_dir / f"docx_{uuid.uuid4().hex}.png"
    tmp_path.write_bytes(image.blob)
    return f'<img src="{tmp_path.as_uri()}"/>'


_HIGHLIGHT_HEX = {
    "AUTO": None,
    "BLACK": "#000000",
    "BLUE": "#0000ff",
    "TURQUOISE": "#00ffff",
    "BRIGHT_GREEN": "#00ff00",
    "PINK": "#ff00ff",
    "RED": "#ff0000",
    "YELLOW": "#ffff00",
    "WHITE": "#ffffff",
    "DARK_BLUE": "#000080",
    "TEAL": "#008080",
    "GREEN": "#008000",
    "VIOLET": "#800080",
    "DARK_RED": "#800000",
    "DARK_YELLOW": "#808000",
    "GRAY_50": "#808080",
    "GRAY_25": "#c0c0c0",
}


def _table_html(table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = html_module.escape(cell.text).replace("\n", "<br/>")
            cells.append(f"<td>{text}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return "<table border='1'>" + "".join(rows) + "</table>"


def load_docx(editor: QTextEdit, path: str | Path) -> None:
    """Carga un documento .docx en el editor como HTML."""
    embedded = _read_embedded_html(path)
    if embedded is not None:
        editor.setHtml(embedded)
        editor.document().setModified(False)
        return
    _load_docx_standard(editor, path)


def _read_embedded_html(path: str | Path) -> str | None:
    """Recupera el HTML completo que rword incrusta al guardar."""
    import zipfile

    with zipfile.ZipFile(path) as archive:
        if "word/rword.html" in archive.namelist():
            return archive.read("word/rword.html").decode("utf-8")
    return None


def _embed_html(path: str | Path, html: str) -> None:
    """Añade el HTML del editor como parte adicional del .docx."""
    import io
    import zipfile

    buffer = io.BytesIO()
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        buffer, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                if "rword.html" not in text:
                    text = text.replace(
                        "</Types>",
                        '<Override PartName="/word/rword.html" '
                        'ContentType="text/html"/></Types>',
                    )
                data = text.encode("utf-8")
            target.writestr(item, data)
        target.writestr("word/rword.html", html.encode("utf-8"))
    Path(path).write_bytes(buffer.getvalue())


def _load_docx_standard(editor: QTextEdit, path: str | Path) -> None:
    """Carga un .docx resolviendo estilos, temas y formato real."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table

    doc = Document(str(path))
    styles = _DocxStyles(doc)
    body = doc.element.body
    parts: list[str] = []
    list_kind: str | None = None

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            is_list, list_type, html = _paragraph_to_html(child, styles, doc)
            if is_list:
                if list_kind != list_type:
                    if list_kind:
                        parts.append("</ol>" if list_kind == "ol" else "</ul>")
                    parts.append("<ol>" if list_type == "ol" else "<ul>")
                    list_kind = list_type
                parts.append(f"<li>{html}</li>")
            else:
                if list_kind:
                    parts.append("</ol>" if list_kind == "ol" else "</ul>")
                    list_kind = None
                parts.append(html)
        elif child.tag == qn("w:tbl"):
            if list_kind:
                parts.append("</ol>" if list_kind == "ol" else "</ul>")
                list_kind = None
            parts.append(_table_html(Table(child, doc)))
    if list_kind:
        parts.append("</ol>" if list_kind == "ol" else "</ul>")

    editor.setHtml("".join(parts))
    default_font, default_size = styles.default_font()
    if default_font or default_size:
        font = editor.font()
        if default_font:
            font.setFamily(default_font)
        if default_size:
            font.setPointSizeF(float(default_size))
        editor.document().setDefaultFont(font)
    editor.document().setModified(False)


def _paragraph_to_html(p_el, styles, doc):
    """Devuelve (es_lista, tipo_lista, html) para un párrafo w:p."""
    from docx.oxml.ns import qn

    content = _runs_html(p_el, styles, doc)
    ppr = p_el.find(qn("w:pPr"))
    style_id = None
    num_pr = None
    if ppr is not None:
        style_el = ppr.find(qn("w:pStyle"))
        if style_el is not None:
            style_id = style_el.get(qn("w:val"))
        num_pr = ppr.find(qn("w:numPr"))

    style_name = styles.style_name(style_id) or ""

    if style_name.startswith("Heading"):
        level = style_name.replace("Heading", "").strip()
        try:
            level = min(6, max(1, int(level)))
        except ValueError:
            level = 1
        return False, None, f"<h{level}>{content}</h{level}>"

    is_list = num_pr is not None or "List Bullet" in style_name or "List Number" in style_name
    if is_list:
        list_type = "ol" if "List Number" in style_name else "ul"
        return True, list_type, content

    props = _paragraph_props(p_el, styles)
    inline = []
    if props.get("align"):
        inline.append(f"text-align:{props['align']}")
    if props.get("left_indent"):
        inline.append(f"margin-left:{props['left_indent']}px")
    if props.get("space_before"):
        inline.append(f"margin-top:{props['space_before']}px")
    if props.get("space_after"):
        inline.append(f"margin-bottom:{props['space_after']}px")
    if props.get("line_spacing"):
        inline.append(f"line-height:{props['line_spacing']:g}")

    if inline:
        return False, None, f"<p style='{'; '.join(inline)}'>{content}</p>"
    return False, None, f"<p>{content}</p>"


def _runs_html(p_el, styles, doc) -> str:
    from docx.oxml.ns import qn

    parts = []
    for child in p_el:
        if child.tag == qn("w:r"):
            rpr = child.find(qn("w:rPr"))
            props = styles.effective_run_props(p_el, child, rpr)
            text = "".join(
                t.text or "" for t in child.findall(qn("w:t"))
            )
            for el in child:
                if el.tag == qn("w:br"):
                    text += "\n"
                elif el.tag == qn("w:tab"):
                    text += "\t"
            blips = child.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if blips:
                embed = blips[0].get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if embed:
                    image_html = _image_html(doc, embed)
                    if image_html:
                        parts.append(image_html)
            parts.append(_text_to_html(text, props))
    return "".join(parts)


def _text_to_html(text: str, props: dict) -> str:
    text = html_module.escape(text)
    if props.get("bold"):
        text = f"<b>{text}</b>"
    if props.get("italic"):
        text = f"<i>{text}</i>"
    if props.get("underline"):
        text = f"<u>{text}</u>"
    if props.get("strike"):
        text = f"<s>{text}</s>"
    if props.get("sup"):
        text = f"<sup>{text}</sup>"
    if props.get("sub"):
        text = f"<sub>{text}</sub>"
    inline = []
    if props.get("color"):
        inline.append(f"color:{props['color']}")
    if props.get("size"):
        inline.append(f"font-size:{props['size']}pt")
    if props.get("font"):
        inline.append(f"font-family:{props['font']}")
    if props.get("highlight"):
        inline.append(f"background-color:{props['highlight']}")
    if inline:
        text = f"<span style='{'; '.join(inline)}'>{text}</span>"
    return text


class _DocxStyles:
    """Resuelve estilos, valores por defecto y fuentes de tema de un .docx."""

    def __init__(self, doc) -> None:
        from docx.oxml.ns import qn

        self.qn = qn
        self.styles: dict[str, dict] = {}
        self.doc_rpr = None
        self.major_font = "Calibri Light"
        self.minor_font = "Calibri"
        self._parse_styles(doc)
        self._parse_theme(doc)

    def _parse_styles(self, doc) -> None:
        styles_element = doc.styles.element
        defaults = styles_element.find(self.qn("w:docDefaults"))
        if defaults is not None:
            rpr_default = defaults.find(self.qn("w:rPrDefault"))
            if rpr_default is not None:
                self.doc_rpr = rpr_default.find(self.qn("w:rPr"))
        for style_el in styles_element.findall(self.qn("w:style")):
            style_id = style_el.get(self.qn("w:styleId"))
            name_el = style_el.find(self.qn("w:name"))
            based = style_el.find(self.qn("w:basedOn"))
            self.styles[style_id] = {
                "name": name_el.get(self.qn("w:val")) if name_el is not None else "",
                "rpr": style_el.find(self.qn("w:rPr")),
                "ppr": style_el.find(self.qn("w:pPr")),
                "based_on": based.get(self.qn("w:val")) if based is not None else None,
            }

    def _parse_theme(self, doc) -> None:
        from lxml import etree

        try:
            parts = doc.part.package.parts
        except Exception:  # noqa: BLE001
            return
        for part in parts:
            try:
                if "theme" not in str(part.partname):
                    continue
            except Exception:  # noqa: BLE001
                continue
            try:
                root = etree.fromstring(part.blob)
            except Exception:  # noqa: BLE001
                continue
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            major = root.find(".//a:majorFont/a:latin", ns)
            minor = root.find(".//a:minorFont/a:latin", ns)
            if major is not None and major.get("typeface"):
                self.major_font = major.get("typeface")
            if minor is not None and minor.get("typeface"):
                self.minor_font = minor.get("typeface")

    def style_name(self, style_id: str | None) -> str | None:
        if not style_id:
            return None
        entry = self.styles.get(style_id)
        return entry["name"] if entry else None

    def _rpr_chain(self, style_id: str | None) -> list:
        chain = []
        seen = set()
        current = style_id
        while current and current not in seen:
            seen.add(current)
            entry = self.styles.get(current)
            if entry is None:
                break
            if entry["rpr"] is not None:
                chain.append(entry["rpr"])
            current = entry["based_on"]
        return chain

    def _ppr_chain(self, style_id: str | None) -> list:
        chain = []
        seen = set()
        current = style_id
        while current and current not in seen:
            seen.add(current)
            entry = self.styles.get(current)
            if entry is None:
                break
            if entry["ppr"] is not None:
                chain.append(entry["ppr"])
            current = entry["based_on"]
        return chain

    def effective_run_props(self, p_el, r_el, rpr) -> dict:
        from docx.oxml.ns import qn

        chain = []
        if rpr is not None:
            chain.append(rpr)
            rstyle = rpr.find(qn("w:rStyle"))
            if rstyle is not None:
                chain.extend(self._rpr_chain(rstyle.get(qn("w:val"))))
        ppr = p_el.find(qn("w:pPr"))
        if ppr is not None:
            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is not None:
                chain.extend(self._rpr_chain(pstyle.get(qn("w:val"))))
        if self.doc_rpr is not None:
            chain.append(self.doc_rpr)
        props = _merge_rpr(chain)
        if "fontTheme" in props:
            theme = props.pop("fontTheme")
            if theme == "majorHAnsi":
                props["font"] = self.major_font
            elif theme == "minorHAnsi":
                props["font"] = self.minor_font
        return props

    def default_font(self) -> tuple[str | None, float | None]:
        props = _merge_rpr([self.doc_rpr])
        font = props.get("font")
        if "fontTheme" in props:
            theme = props["fontTheme"]
            font = self.major_font if theme == "majorHAnsi" else self.minor_font
        return font, props.get("size")


def _merge_rpr(chain) -> dict:
    props: dict = {}
    for rpr in chain:
        if rpr is None:
            continue
        for key, value in _rpr_props(rpr).items():
            if key not in props:
                props[key] = value
    return props


def _rpr_props(rpr) -> dict:
    from docx.oxml.ns import qn

    props: dict = {}
    b = rpr.find(qn("w:b"))
    if b is not None:
        props["bold"] = b.get(qn("w:val")) not in ("0", "false", "off")
    i = rpr.find(qn("w:i"))
    if i is not None:
        props["italic"] = i.get(qn("w:val")) not in ("0", "false", "off")
    if rpr.find(qn("w:u")) is not None:
        props["underline"] = True
    if rpr.find(qn("w:strike")) is not None:
        props["strike"] = True
    va = rpr.find(qn("w:vertAlign"))
    if va is not None:
        value = va.get(qn("w:val"))
        if value == "superscript":
            props["sup"] = True
        elif value == "subscript":
            props["sub"] = True
    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        import contextlib

        with contextlib.suppress(TypeError, ValueError):
            props["size"] = int(sz.get(qn("w:val"))) / 2.0
    color = rpr.find(qn("w:color"))
    if color is not None:
        value = color.get(qn("w:val"))
        if value and value != "auto":
            props["color"] = ("#" + value).lower()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is not None:
        theme = fonts.get(qn("w:asciiTheme"))
        if theme:
            props["fontTheme"] = theme
        else:
            name = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
            if name:
                props["font"] = name
    highlight = rpr.find(qn("w:highlight"))
    if highlight is not None:
        value = highlight.get(qn("w:val"))
        color = _HIGHLIGHT_HEX.get((value or "").upper())
        if color:
            props["highlight"] = color
    return props


def _paragraph_props(p_el, styles) -> dict:
    from docx.oxml.ns import qn

    ppr = p_el.find(qn("w:pPr"))
    chain = []
    if ppr is not None:
        chain.append(ppr)
        pstyle = ppr.find(qn("w:pStyle"))
        if pstyle is not None:
            chain.extend(styles._ppr_chain(pstyle.get(qn("w:val"))))

    result: dict = {}
    for element in chain:
        if element is None:
            continue
        if "align" not in result:
            jc = element.find(qn("w:jc"))
            if jc is not None:
                result["align"] = {
                    "start": "left", "end": "right", "left": "left",
                    "center": "center", "right": "right",
                    "justify": "justify", "both": "justify",
                    "distribute": "justify",
                }.get(jc.get(qn("w:val")))
        ind = element.find(qn("w:ind"))
        if ind is not None and "left_indent" not in result:
            left = ind.get(qn("w:left"))
            if left:
                result["left_indent"] = int(round(int(left) / 567.0))
        spacing = element.find(qn("w:spacing"))
        if spacing is not None:
            if "space_before" not in result and spacing.get(qn("w:before")):
                result["space_before"] = int(round(int(spacing.get(qn("w:before"))) / 567.0))
            if "space_after" not in result and spacing.get(qn("w:after")):
                result["space_after"] = int(round(int(spacing.get(qn("w:after"))) / 567.0))
            if "line_spacing" not in result:
                line = spacing.get(qn("w:line"))
                rule = spacing.get(qn("w:lineRule"))
                if line and rule == "auto":
                    result["line_spacing"] = int(line) / 240.0
    return result
def save_docx(editor: QTextEdit, path: str | Path) -> None:
    """Guarda el contenido del editor en un documento .docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    document = editor.document()
    _write_blocks(doc, document, WD_ALIGN_PARAGRAPH)
    doc.save(str(path))
    _embed_html(path, editor.toHtml())


def _write_blocks(doc, document: QTextDocument, wd_align) -> None:

    block = document.begin()
    while block.isValid():
        frame = block.textList()
        table = _table_of_block(document, block)
        if table is not None:
            _write_table(doc, table, wd_align)
            block = table.lastCursorPosition().block()
        elif frame is not None:
            style = "List Bullet" if frame.format().style() in (
                QTextListFormat.Style.ListDisc,
                QTextListFormat.Style.ListCircle,
                QTextListFormat.Style.ListSquare,
            ) else "List Number"
            paragraph = doc.add_paragraph(style=style)
            _fill_paragraph(paragraph, block, document)
        else:
            paragraph = doc.add_paragraph()
            _fill_paragraph(paragraph, block, document)
        block = block.next()


def _table_of_block(document: QTextDocument, block):
    position = block.position()
    for frame in document.rootFrame().childFrames():
        if (
            isinstance(frame, QTextTable)
            and frame.firstPosition() <= position <= frame.lastPosition()
        ):
            return frame
    return None


def _fill_paragraph(paragraph, block, document) -> None:
    from PySide6.QtCore import Qt

    fmt = block.blockFormat()
    alignment = fmt.alignment()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if alignment & Qt.AlignmentFlag.AlignHCenter:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment & Qt.AlignmentFlag.AlignRight:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment & Qt.AlignmentFlag.AlignJustify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if fmt.leftMargin():
        paragraph.paragraph_format.left_indent = int(fmt.leftMargin() / 1.5)
    if fmt.topMargin() or fmt.bottomMargin():
        paragraph.paragraph_format.space_before = int(fmt.topMargin() / 1.5)
        paragraph.paragraph_format.space_after = int(fmt.bottomMargin() / 1.5)

    iterator = block.begin()
    while not iterator.atEnd():
        fragment = iterator.fragment()
        text = fragment.text()
        fmt = fragment.charFormat()
        if fmt.isImageFormat():
            _append_image(paragraph, fmt, document)
        elif text:
            run = paragraph.add_run(text)
            if fmt.fontWeight() >= 700:
                run.bold = True
            if fmt.fontItalic():
                run.italic = True
            if fmt.fontUnderline():
                run.underline = True
            if fmt.fontStrikeOut():
                run.font.strike = True
            if fmt.fontPointSize() > 0:
                run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(
                    fmt.fontPointSize()
                )
            families = fmt.fontFamilies()
            if families and families[0]:
                run.font.name = families[0]
            highlight = fmt.background().color()
            if highlight.isValid() and highlight.name() not in (
                "#ffffff", "#00000000", "#000000", "#ff000000",
            ):
                index = _nearest_highlight(highlight.name())
                if index is not None:
                    run.font.highlight_color = index
            vertical = fmt.verticalAlignment()
            from PySide6.QtGui import QTextCharFormat

            if vertical == QTextCharFormat.VerticalAlignment.AlignSuperScript:
                run.font.superscript = True
            elif vertical == QTextCharFormat.VerticalAlignment.AlignSubScript:
                run.font.subscript = True
            foreground = fmt.foreground().color()
            if foreground.isValid() and foreground.name() != "#000000":
                run.font.color.rgb = __import__(
                    "docx.shared", fromlist=["RGBColor"]
                ).RGBColor(*foreground.getRgb()[:3])
        iterator += 1


def _nearest_highlight(hex_color: str):
    from docx.enum.text import WD_COLOR_INDEX

    try:
        target = tuple(int(hex_color[i:i + 2], 16) for i in (1, 3, 5))
    except ValueError:
        return None
    best = None
    best_dist = float("inf")
    for name, value in _HIGHLIGHT_HEX.items():
        if value is None:
            continue
        color = tuple(int(value[i:i + 2], 16) for i in (1, 3, 5))
        dist = sum((a - b) ** 2 for a, b in zip(target, color, strict=True))
        if dist < best_dist:
            best_dist = dist
            best = name
    return getattr(WD_COLOR_INDEX, best, None)


def _append_image(paragraph, char_format, document) -> None:
    import uuid

    from PySide6.QtGui import QImage, QTextFormat

    name = char_format.stringProperty(QTextFormat.Property.ImageName)
    if not name:
        return
    resource = document.resource(QTextDocument.ImageResource, QUrl(name))
    if resource is None or not isinstance(resource, QImage) or resource.isNull():
        return
    temp_dir = Path.home() / ".cache" / "rword"
    temp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = temp_dir / f"rword_{uuid.uuid4().hex}.png"
    resource.save(str(tmp_path))
    try:
        paragraph.add_run().add_picture(str(tmp_path))
    finally:
        tmp_path.unlink(missing_ok=True)


def _write_table(doc, table, wd_align) -> None:
    rows = table.rows()
    cols = table.columns()
    if rows == 0 or cols == 0:
        return
    doc_table = doc.add_table(rows=rows, cols=cols)
    doc_table.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            cell = table.cellAt(r, c)
            text = cell.firstCursorPosition().block().text()
            doc_table.cell(r, c).text = text
    doc.add_paragraph()
